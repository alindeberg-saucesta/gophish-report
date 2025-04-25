#!/usr/bin/env python3

import requests
import os
from flask import (Flask, request, jsonify, render_template_string, send_file,
                   session, redirect, url_for, flash)
from functools import wraps
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl # For Excel
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
from io import BytesIO
import json
from dateutil import parser # Using dateutil for more robust parsing
from dateutil.tz import UTC # To handle timezone explicitly
import traceback # For detailed error logging
from datetime import datetime, date # For type checking in Excel date formatting

# --- Configuration ---
DEFAULT_GOPHISH_URL = os.getenv('GOPHISH_URL', "https://<your url>:3333") # Use ENV var or default
DEFAULT_GOPHISH_API_KEY = os.getenv('GOPHISH_API_KEY', "<your GoPhish API Key>") # Use ENV var or default

# --- Simple Login Credentials (INSECURE - FOR DEMO ONLY) ---
USERNAME = "admin"
PASSWORD = "yourpassword" # Use proper hashing in production

# --- Create the Flask App Instance ---
app = Flask(__name__)

# --- Set Secret Key for Sessions (CHANGE THIS AND KEEP SECRET!) ---
# Use a long, random string. Best practice: Load from environment variable.
app.secret_key = os.getenv('SECRET_KEY', '!!change-this-to-a-real-secret-key-39z56y!!') # <-- CHANGE THIS!!!

# --- Helper Functions ---

def get_gophish_headers(api_key):
    """Creates the authorization header."""
    return {'Authorization': f'Bearer {api_key}'}

def format_datetime(dt_string):
    """Formats ISO8601 datetime string to a more readable format."""
    if not dt_string or dt_string.startswith("0001"): return "N/A"
    try:
        dt_obj = parser.isoparse(dt_string)
        return dt_obj.strftime('%B %d %Y %I:%M:%S %p %Z').replace(" 0", " ").replace(" AM", "am").replace(" PM", "pm")
    except (ValueError, TypeError) as e:
        app.logger.error(f"Error parsing date: {dt_string} - {e}")
        return dt_string

def parse_timeline_details(details_str):
    """Parses details JSON, extracts payload summary, browser info, and IP."""
    if not details_str: return None, None, None # payload, browser, ip
    try:
        details = json.loads(details_str)
        payload = details.get('payload', {})
        browser = details.get('browser', {})
        browser_info_str = None; user_agent = browser.get('user-agent')
        if user_agent:
            os_part, browser_part = "Unknown OS", "Unknown Browser"
            if "Mac OS" in user_agent: os_part = "Mac OS"
            elif "Windows" in user_agent: os_part = "Windows"
            elif "Linux" in user_agent: os_part = "Linux"
            elif "Android" in user_agent: os_part = "Android"
            elif "iPhone" in user_agent or "iPad" in user_agent: os_part = "iOS"
            if "Edg/" in user_agent: browser_part = "Edge"
            elif "Chrome" in user_agent: browser_part = "Chrome"
            elif "Firefox" in user_agent: browser_part = "Firefox"
            elif "Safari" in user_agent and "Chrome" not in user_agent: browser_part = "Safari"
            browser_info_str = f"{os_part} | {browser_part}"
        ip_address = browser.get('address')
        payload_info = None
        if payload:
            if 'password' in payload or 'username' in payload: payload_info = "Credentials Submitted"
            else:
                payload_summary = ", ".join([f"{k}: {v}" for k, v in payload.items() if k != 'rid'])
                if payload_summary: payload_info = f"Data: {payload_summary}"
        return payload_info, browser_info_str, ip_address
    except json.JSONDecodeError: return "Invalid JSON", None, None
    except Exception as e:
        app.logger.error(f"Error parsing details: {e}\n{traceback.format_exc()}")
        return "Error parsing details", None, None

# --- Login Required Decorator ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session:
            flash('Please log in to access this page.', 'warning')
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

# --- Flask Routes ---

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        submitted_user = request.form.get('username')
        submitted_pass = request.form.get('password')
        if submitted_user == USERNAME and submitted_pass == PASSWORD:
            session['logged_in'] = True; session['username'] = submitted_user
            flash('Login successful!', 'success')
            next_url = request.args.get('next')
            if next_url and (next_url.startswith('/') or next_url.startswith(request.host_url)): return redirect(next_url)
            else: return redirect(url_for('index'))
        else:
            flash('Invalid username or password.', 'danger')
            return render_template_string(LOGIN_TEMPLATE)
    if 'logged_in' in session: return redirect(url_for('index'))
    return render_template_string(LOGIN_TEMPLATE)

@app.route('/logout')
def logout():
    session.pop('logged_in', None); session.pop('username', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    try:
        with open('index.html', 'r') as f: html_content = f.read()
        gophish_url = os.getenv('GOPHISH_URL', DEFAULT_GOPHISH_URL)
        gophish_api_key = os.getenv('GOPHISH_API_KEY', DEFAULT_GOPHISH_API_KEY)
        return render_template_string(html_content,
                                      gophish_url=gophish_url,
                                      gophish_api_key=gophish_api_key,
                                      logged_in=session.get('logged_in', False))
    except FileNotFoundError:
        app.logger.error("index.html not found.")
        return "Error: index.html not found.", 500
    except Exception as e:
        app.logger.error(f"Error serving index.html: {e}\n{traceback.format_exc()}")
        return "An internal server error occurred.", 500

@app.route('/api/campaigns', methods=['GET'])
@login_required
def get_campaigns():
    gophish_url = request.args.get('gophish_url', os.getenv('GOPHISH_URL', DEFAULT_GOPHISH_URL))
    api_key = request.args.get('api_key', os.getenv('GOPHISH_API_KEY', DEFAULT_GOPHISH_API_KEY))
    if not gophish_url or not api_key: return jsonify({"error": "Gophish URL and API Key are required"}), 400
    api_endpoint = f"{gophish_url.rstrip('/')}/api/campaigns/"; headers = get_gophish_headers(api_key); response = None
    try:
        # !! verify=False is INSECURE !!
        response = requests.get(api_endpoint, headers=headers, verify=False, timeout=10)
        response.raise_for_status()
        campaigns = response.json()
        campaign_list = [{"id": c.get("id"), "name": c.get("name")} for c in campaigns if c.get("id") and c.get("name")]
        return jsonify(campaign_list)
    except requests.exceptions.Timeout: return jsonify({"error": "Connection to Gophish API timed out."}), 504
    except requests.exceptions.ConnectionError: return jsonify({"error": "Could not connect to Gophish API. Check URL/network."}), 502
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code; error_text = e.response.text[:200]
        app.logger.error(f"Gophish API Error {status_code} from {api_endpoint}: {error_text}")
        if status_code == 401: return jsonify({"error": "Gophish API Authentication Failed (401). Check API Key."}), 401
        else: return jsonify({"error": f"Gophish API returned error {status_code}. {error_text}"}), status_code
    except Exception as e:
        app.logger.error(f"Unexpected error fetching campaigns: {e}\n{traceback.format_exc()}")
        return jsonify({"error": "An unexpected server error occurred."}), 500

@app.route('/api/generate-report/<int:campaign_id>', methods=['GET'])
@login_required
def generate_report(campaign_id):
    # (Word Report Generation - Code remains the same as previous full version)
    gophish_url = request.args.get('gophish_url', os.getenv('GOPHISH_URL', DEFAULT_GOPHISH_URL))
    api_key = request.args.get('api_key', os.getenv('GOPHISH_API_KEY', DEFAULT_GOPHISH_API_KEY))
    if not gophish_url or not api_key: return jsonify({"error": "Gophish URL and API Key are required"}), 400
    headers = get_gophish_headers(api_key)
    campaign_endpoint = f"{gophish_url.rstrip('/')}/api/campaigns/{campaign_id}"
    summary_endpoint = f"{gophish_url.rstrip('/')}/api/campaigns/{campaign_id}/summary"
    campaign_response = None; summary_response = None
    try:
        campaign_response = requests.get(campaign_endpoint, headers=headers, verify=False, timeout=15) # !! verify=False !!
        campaign_response.raise_for_status(); campaign_data = campaign_response.json()
        summary_response = requests.get(summary_endpoint, headers=headers, verify=False, timeout=10) # !! verify=False !!
        summary_response.raise_for_status(); summary_data = summary_response.json(); stats = summary_data.get('stats', {})
        document = Document()
        title = document.add_heading(f"Gophish Campaign Report: {campaign_data.get('name', 'N/A')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; document.add_paragraph()
        document.add_heading('Campaign Statistics', level=2)
        stats_list = [f"Total Recipients:\t{stats.get('total', 0)}", f"Emails Sent:\t\t{stats.get('sent', 0)}", f"Emails Opened:\t{stats.get('opened', 0)}", f"Clicked Link:\t\t{stats.get('clicked', 0)}", f"Submitted Data:\t{stats.get('submitted_data', 0)}", f"Email Reported:\t{stats.get('email_reported', 0)}", f"Errors:\t\t{stats.get('error', 0)}"]
        for item in stats_list: document.add_paragraph(item)
        document.add_paragraph(); document.add_heading('Details', level=2)
        results = campaign_data.get('results', []); timeline = campaign_data.get('timeline', [])
        if not results: document.add_paragraph("No results found.")
        else:
            results.sort(key=lambda x: x.get('email', ''))
            for result in results:
                user_email = result.get('email', 'N/A'); user_first = result.get('first_name', ''); user_last = result.get('last_name', ''); user_name = f"{user_first} {user_last}".strip() if user_first or user_last else user_email; result_id = result.get('id', 'N/A')
                user_p = document.add_paragraph(); user_p.add_run(f"{user_name}").bold = True; user_p.add_run(f" ({user_email})"); user_p.add_run(f"\tStatus: "); user_p.add_run(f"{result.get('status', 'N/A')}").italic = True
                if result.get('reported'): user_p.add_run("\t(Reported)").bold = True
                position = result.get('position')
                if position: pos_p = document.add_paragraph(); pos_p.paragraph_format.left_indent = Inches(0.25); run = pos_p.add_run(f"Position: {position}"); run.font.size = Pt(9); run.italic = True
                timeline_heading = document.add_paragraph(); timeline_heading.paragraph_format.left_indent = Inches(0.25); run = timeline_heading.add_run(f"Timeline (Result ID: {result_id})"); run.font.size = Pt(10); run.italic = True
                user_timeline = [event for event in timeline if event.get('email') == user_email]; created_event = next((event for event in timeline if event.get('message') == 'Campaign Created'), None)
                display_timeline = [];
                if created_event: display_timeline.append(created_event)
                display_timeline.extend(user_timeline); display_timeline = sorted(list({event['time']: event for event in display_timeline}.values()), key=lambda x: parser.isoparse(x['time']))
                if not user_timeline: p = document.add_paragraph(); p.paragraph_format.left_indent = Inches(0.5); p.add_run("  No specific timeline events recorded.").italic = True
                else:
                    for event in display_timeline:
                        event_time_str = format_datetime(event.get('time')); event_message = event.get('message', 'N/A')
                        p = document.add_paragraph(); p.paragraph_format.left_indent = Inches(0.5); p.add_run(f"{event_message}: ").bold = True; p.add_run(event_time_str)
                        payload_details, browser_details, ip_address = parse_timeline_details(event.get('details'))
                        if browser_details: detail_p = document.add_paragraph(); detail_p.paragraph_format.left_indent = Inches(0.75); run = detail_p.add_run(f"Context: {browser_details}{f' | IP: {ip_address}' if ip_address else ''}"); run.font.size = Pt(9); run.italic = True
                        if payload_details: detail_p = document.add_paragraph(); detail_p.paragraph_format.left_indent = Inches(0.75); run = detail_p.add_run(f"Payload: {payload_details}"); run.font.size = Pt(9); run.italic = True
                document.add_paragraph()
        file_stream = BytesIO(); document.save(file_stream); file_stream.seek(0)
        campaign_name = campaign_data.get('name', 'campaign'); filename_safe_name = "".join(c if c.isalnum() else "_" for c in campaign_name); filename = f"Gophish_Report_{campaign_id}_{filename_safe_name}.docx"
        return send_file(file_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except requests.exceptions.Timeout: return jsonify({"error": "Gophish API timed out generating Word report."}), 504
    except requests.exceptions.ConnectionError: return jsonify({"error": "Could not connect to Gophish API generating Word report."}), 502
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code; error_text = e.response.text[:200]; app.logger.error(f"Gophish API Error {status_code} for Word report {campaign_id}: {error_text}")
        if status_code == 404: return jsonify({"error": f"Campaign ID {campaign_id} not found (404)."}), 404
        return jsonify({"error": f"Gophish API returned error {status_code} generating Word report. {error_text}"}), status_code
    except Exception as e:
        app.logger.error(f"Unexpected error generating Word report {campaign_id}: {e}\n{traceback.format_exc()}")
        return jsonify({"error": "An unexpected server error occurred generating Word report."}), 500


@app.route('/api/generate-excel-report/<int:campaign_id>', methods=['GET'])
@login_required
def generate_excel_report(campaign_id):
    """API endpoint to generate and return an Excel (.xlsx) report."""
    gophish_url = request.args.get('gophish_url', os.getenv('GOPHISH_URL', DEFAULT_GOPHISH_URL))
    api_key = request.args.get('api_key', os.getenv('GOPHISH_API_KEY', DEFAULT_GOPHISH_API_KEY))
    if not gophish_url or not api_key: return jsonify({"error": "Gophish URL and API Key are required"}), 400

    headers = get_gophish_headers(api_key)
    campaign_endpoint = f"{gophish_url.rstrip('/')}/api/campaigns/{campaign_id}"
    summary_endpoint = f"{gophish_url.rstrip('/')}/api/campaigns/{campaign_id}/summary"
    campaign_response = None; summary_response = None
    try:
        # Fetch Campaign & Summary Data
        campaign_response = requests.get(campaign_endpoint, headers=headers, verify=False, timeout=15) # !! verify=False !!
        campaign_response.raise_for_status(); campaign_data = campaign_response.json()
        summary_response = requests.get(summary_endpoint, headers=headers, verify=False, timeout=10) # !! verify=False !!
        summary_response.raise_for_status(); summary_data = summary_response.json(); stats = summary_data.get('stats', {})

        # Generate Excel Workbook
        wb = openpyxl.Workbook(); ws_summary = wb.active; ws_summary.title = "Summary"
        header_font = Font(bold=True); title_font = Font(bold=True, size=14)
        ws_summary['A1'] = f"Gophish Campaign Report: {campaign_data.get('name', 'N/A')}"
        ws_summary['A1'].font = title_font; ws_summary.merge_cells('A1:D1')
        summary_info = {"ID": campaign_data.get('id'), "Name": campaign_data.get('name'), "Status": campaign_data.get('status'), "Created Date": format_datetime(campaign_data.get('created_date')), "Launch Date": format_datetime(campaign_data.get('launch_date')), "Completed Date": format_datetime(campaign_data.get('completed_date')), "Template": campaign_data.get('template', {}).get('name', 'N/A'), "Landing Page": campaign_data.get('page', {}).get('name', 'N/A'), "URL": campaign_data.get('url')}
        row = 3
        for key, value in summary_info.items(): ws_summary[f'A{row}'] = key; ws_summary[f'A{row}'].font = header_font; ws_summary[f'B{row}'] = str(value) if value is not None else 'N/A'; row += 1
        row += 1; ws_summary[f'A{row}'] = "Campaign Statistics"; ws_summary[f'A{row}'].font = header_font; row += 1
        stats_info = {"Total Recipients": stats.get('total', 0), "Emails Sent": stats.get('sent', 0), "Emails Opened": stats.get('opened', 0), "Clicked Link": stats.get('clicked', 0), "Submitted Data": stats.get('submitted_data', 0), "Email Reported": stats.get('email_reported', 0), "Errors": stats.get('error', 0)}
        for key, value in stats_info.items(): ws_summary[f'A{row}'] = key; ws_summary[f'B{row}'] = value; row += 1
        ws_summary.column_dimensions['A'].width = 25; ws_summary.column_dimensions['B'].width = 40

        ws_details = wb.create_sheet(title="Detailed Timeline")
        headers = ["Email", "First Name", "Last Name", "Position", "Final Status", "Reported", "Event Time (UTC)", "Event Message", "Browser/OS", "IP Address", "Payload Data"] # Clarified Timezone
        for col_idx, header_title in enumerate(headers, 1): cell = ws_details.cell(row=1, column=col_idx, value=header_title); cell.font = header_font
        user_details_map = {};
        for r in campaign_data.get('results', []):
             email = r.get('email');
             if email: user_details_map[email] = {"first_name": r.get('first_name', ''), "last_name": r.get('last_name', ''), "position": r.get('position', ''), "status": r.get('status', ''), "reported": r.get('reported', False)}
        row = 2; timeline = sorted(campaign_data.get('timeline', []), key=lambda x: parser.isoparse(x['time']))
        for event in timeline:
            event_email = event.get('email'); user_info = user_details_map.get(event_email, {})
            payload_info, browser_info, ip_address = parse_timeline_details(event.get('details'))

            # --- FIX for Excel Timezone ---
            excel_event_time = None # This will hold the value written to Excel
            event_time_str = event.get('time')
            if event_time_str:
                try:
                    aware_dt = parser.isoparse(event_time_str)
                    excel_event_time = aware_dt.replace(tzinfo=None) # Convert to naive datetime
                except (ValueError, TypeError):
                    excel_event_time = event_time_str # Fallback to string
            else:
                 excel_event_time = "N/A"
            # --- End of FIX ---

            data_row = [event_email or "(Campaign Event)", user_info.get('first_name', ''), user_info.get('last_name', ''), user_info.get('position', ''), user_info.get('status', '') if event_email else '', 'Yes' if user_info.get('reported') else 'No' if event_email else '', excel_event_time, event.get('message', 'N/A'), browser_info, ip_address, payload_info]
            for col_idx, value in enumerate(data_row, 1):
                 cell = ws_details.cell(row=row, column=col_idx, value=value)
                 if isinstance(value, (datetime, date)): cell.number_format = 'yyyy-mm-dd hh:mm:ss'
                 elif col_idx == 7 and not isinstance(value, (datetime, date)): cell.number_format = '@' # Format fallback string as Text
            row += 1
        for i, col_letter in enumerate(get_column_letter(idx) for idx in range(1, len(headers) + 1)): ws_details.column_dimensions[col_letter].width = max(len(headers[i]) + 2, 20)

        file_stream = BytesIO(); wb.save(file_stream); file_stream.seek(0) # Save to memory
        campaign_name = campaign_data.get('name', 'campaign'); filename_safe_name = "".join(c if c.isalnum() else "_" for c in campaign_name); filename = f"Gophish_Report_{campaign_id}_{filename_safe_name}.xlsx"
        return send_file(file_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except requests.exceptions.Timeout: return jsonify({"error": "Gophish API timed out generating Excel report."}), 504
    except requests.exceptions.ConnectionError: return jsonify({"error": "Could not connect to Gophish API generating Excel report."}), 502
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code; error_text = e.response.text[:200]; app.logger.error(f"Gophish API Error {status_code} for Excel report {campaign_id}: {error_text}")
        if status_code == 404: return jsonify({"error": f"Campaign ID {campaign_id} not found (404)."}), 404
        return jsonify({"error": f"Gophish API returned error {status_code} generating Excel report. {error_text}"}), status_code
    except Exception as e:
        app.logger.error(f"Unexpected error generating Excel report {campaign_id}: {e}\n{traceback.format_exc()}")
        return jsonify({"error": "An unexpected server error occurred generating Excel report."}), 500


# --- Login Template Definition ---
LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8"> <meta name="viewport" content="width=device-width, initial-scale=1.0"> <title>Login - Gophish Reporter</title>
    <style> body { font-family: sans-serif; background-color: #f4f4f4; display: flex; justify-content: center; align-items: center; height: 100vh; margin: 0;} .login-container { background-color: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); width: 300px; } h2 { text-align: center; margin-bottom: 20px; color: #333; } label { display: block; margin-bottom: 5px; color: #555; } input[type="text"], input[type="password"] { width: 100%; padding: 10px; margin-bottom: 15px; border: 1px solid #ccc; border-radius: 4px; box-sizing: border-box; } button { width: 100%; padding: 10px; background-color: #5cb85c; color: white; border: none; border-radius: 4px; cursor: pointer; font-size: 16px; } button:hover { background-color: #4cae4c; } .flash-messages { list-style: none; padding: 0; margin-bottom: 15px; } .flash-messages li { padding: 10px; border-radius: 4px; margin-bottom: 10px; font-size: 0.9em; } .flash-messages .success { background-color: #dff0d8; color: #3c763d; border: 1px solid #d6e9c6; } .flash-messages .danger { background-color: #f2dede; color: #a94442; border: 1px solid #ebccd1; } .flash-messages .warning { background-color: #fcf8e3; color: #8a6d3b; border: 1px solid #faebcc; } .flash-messages .info { background-color: #d9edf7; color: #31708f; border: 1px solid #bce8f1; } </style>
</head>
<body>
    <div class="login-container"> <h2>Login</h2> {% with messages = get_flashed_messages(with_categories=true) %} {% if messages %} <ul class="flash-messages"> {% for category, message in messages %} <li class="{{ category }}">{{ message }}</li> {% endfor %} </ul> {% endif %} {% endwith %} <form method="POST"> <div><label for="username">Username:</label><input type="text" id="username" name="username" required></div> <div><label for="password">Password:</label><input type="password" id="password" name="password" required></div> <button type="submit">Login</button> </form> </div>
</body>
</html>
"""

# --- Run the App ---
if __name__ == '__main__':
    # !! Set debug=False for production !!
    # Consider using a production WSGI server (like Gunicorn or Waitress) instead of app.run()
    app.run(host='0.0.0.0', port=5000, debug=True)
